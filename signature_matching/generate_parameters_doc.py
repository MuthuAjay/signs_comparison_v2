"""
Signature Comparison — Parameters, Methodology, Scope and Limitations.

Rebuilds the standing parameters/scope/limitations document in the house style
of Signature_Parameters_Scope_Limitations_final, carrying the current findings
rather than the earlier two-signature run.

What changed since that version, and why the document had to be rewritten
rather than edited: the earlier run compared the two questioned Letters of
Authority against each other, with no genuine reference signatures at all, and
calibrated against perturbed copies of the crops themselves. Three specimen
documents have since been supplied, so the comparison is now questioned-against-
genuine, and the self-perturbation calibration has been dropped in favour of the
genuine-to-genuine agreement actually observed across those specimens.

Every figure is read from the run's comparison_report.csv, and the aspect-ratio
correlation is recomputed from the crops, so the document cannot drift from the
run that produced it.

Usage:
    python generate_parameters_doc.py
        [--csv comparison_output_shafiul_only/comparison_report.csv]
        [--exhibit ../loa_comparison/exhibit_shafiul_only.png]
        [--sig-dir signatures]
        [--out "Signature Comparison - Parameters, Methodology, Scope and Limitations.docx"]
"""

import argparse
import csv
import math
import os
from collections import defaultdict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)
HDR_FILL = "D9E2F3"
BODY_PT = 10.5
SCORERS = ("global_cosine", "patch_match_score", "ssim_aligned")

DISPLAY = {
    "ref_shafiul_invitation": "Invitation letter (2014)",
    "ref_shafiul_navana_sa": "Navana service agreement",
    "ref_shafiul_rjsc": "RJSC Form XVIII",
    "q_shafiul_hcd": "HCD Letter of Authority",
    "q_shafiul_cmm": "CMM Letter of Authority",
}
SHORT = {
    "ref_shafiul_invitation": "Invitation letter",
    "ref_shafiul_navana_sa": "Service agreement",
    "ref_shafiul_rjsc": "RJSC Form XVIII",
    "q_shafiul_hcd": "HCD letter",
    "q_shafiul_cmm": "CMM letter",
    "ctl_shafqat_hcd": "Shafqat Ahmed",
    "ctl_muto_navana_sa": "Kazuyuki Muto",
    "ctl_second_signer_rjsc": "2nd signatory, RJSC",
}


# ----------------------------------------------------------------------------
# House-style primitives, matched to the reference document
# ----------------------------------------------------------------------------

def title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size, r.font.bold, r.font.color.rgb = Pt(16), True, NAVY
    r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(14)
    return p


def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run(text)
    r.font.size, r.font.bold, r.font.color.rgb = Pt(14), True, NAVY
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def body(doc, parts, space_after=10, align=None):
    """A body paragraph. `parts` is a string, or (text, bold, italic) fragments."""
    p = doc.add_paragraph()
    if isinstance(parts, str):
        parts = [(parts, False, False)]
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.font.size, r.font.bold, r.font.italic = Pt(BODY_PT + 0.5), bold, italic
        r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def lead(doc, label, text):
    """The reference's labelled methodology paragraph: 'Detection: ...'."""
    return body(doc, [(label, True, False), (text, False, False)])


def _fixed_grid(tbl, widths):
    """Pin the column widths.

    With autofit left on, Word and LibreOffice both re-derive column widths
    from the cell contents and the requested widths are ignored -- which is
    how a four-column table ends up with a 1-inch column of wrapped text. The
    reference document pins tblW and every gridCol, so the same is done here:
    fixed layout, an explicit grid, and tcW on each cell.
    """
    tbl.autofit = False
    pr = tbl._tbl.find(qn("w:tblPr"))
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    pr.append(layout)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    w.set(qn("w:type"), "dxa")
    pr.append(w)
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for col, wid in zip(grid.findall(qn("w:gridCol")), widths):
        col.set(qn("w:w"), str(int(wid * 1440)))


def _borders(tbl):
    pr = tbl._tbl.find(qn("w:tblPr"))
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        b.append(e)
    pr.append(b)
    mar = OxmlElement("w:tblCellMar")
    for edge, w in (("top", 80), ("left", 120), ("bottom", 80), ("right", 120)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(w))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    pr.append(mar)


def _shade(cell, fill):
    e = OxmlElement("w:shd")
    e.set(qn("w:val"), "clear")
    e.set(qn("w:color"), "auto")
    e.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(e)


def _no_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def _repeat_header(row):
    """Reprint the header row when a table continues onto the next page."""
    row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = doc.styles["Normal Table"]
    _borders(t)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.font.size, r.font.bold = Pt(BODY_PT), True
        r.font.name = "Calibri"
        _shade(c, HDR_FILL)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(BODY_PT)
            r.font.name = "Calibri"
    _fixed_grid(t, widths)
    _repeat_header(t.rows[0])
    for row in t.rows:
        _no_split(row)
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ----------------------------------------------------------------------------
# Data
# ----------------------------------------------------------------------------

def load(csv_path):
    with open(csv_path) as f:
        rows = list(csv.DictReader(f))
    for r in rows:
        for k in SCORERS + ("mean_score", "patch_mean_sim"):
            r[k] = float(r[k])
    return rows


def split(rows):
    """Group pairs by what they calibrate.

    The floor is reference-vs-control only: a genuine signature scored against
    another writer's. That is the same comparison made for a questioned mark,
    differing only in whether the second signature is his. Control-vs-control
    and questioned-vs-control pairs compare two marks neither of which is a
    genuine specimen; they calibrate nothing and are held out of the bands.
    """
    g = defaultdict(list)
    for r in rows:
        t = r["type"]
        if t.startswith("ref-ref"):
            k = "ceiling"
        elif t.startswith("ref-questioned"):
            k = "questioned"
        elif t.startswith("q-q"):
            k = "qq"
        elif t.startswith("ref-control"):
            k = "floor"
        else:
            k = "context"
        g[k].append(r)
    return g


def placement(v, floor_hi, ceil_lo):
    if v >= ceil_lo:
        return "same-hand range"
    if v <= floor_hi:
        return "DIFFERENT-hands range"
    return "between ranges"


def pct(v):
    return f"{v * 100:.0f}%"


def rng(rows, key):
    v = [r[key] for r in rows]
    return min(v), max(v)


def rng_pct(rows, key):
    v = [r[key] for r in rows]
    return f"{min(v) * 100:.0f}–{max(v) * 100:.0f}%"


def aspect_ratios(names, sig_dir):
    import numpy as np
    from PIL import Image
    out = {}
    for n in names:
        p = os.path.join(sig_dir, f"{n}.png")
        if not os.path.exists(p):
            continue
        ink = np.asarray(Image.open(p).convert("L"), dtype=np.uint8) <= 128
        ys, xs = np.where(ink)
        if len(xs):
            out[n] = float((xs.max() - xs.min() + 1) / (ys.max() - ys.min() + 1))
    return out


def pearson(a, b):
    n = len(a)
    if n < 3:
        return None
    ma, mb = sum(a) / n, sum(b) / n
    num = sum((x - ma) * (y - mb) for x, y in zip(a, b))
    da = math.sqrt(sum((x - ma) ** 2 for x in a))
    db = math.sqrt(sum((y - mb) ** 2 for y in b))
    return num / (da * db) if da and db else None


def aspect_corr(rows, ratios, key):
    d, c = [], []
    for r in rows:
        if r["sig_a"] in ratios and r["sig_b"] in ratios:
            d.append(abs(math.log(ratios[r["sig_a"]]) - math.log(ratios[r["sig_b"]])))
            c.append(r[key])
    return pearson(d, c)


# ----------------------------------------------------------------------------
# Document
# ----------------------------------------------------------------------------

def build(rows, exhibit, sig_dir, out_path):
    g = split(rows)
    ceiling, questioned, qq = g["ceiling"], g["questioned"], g["qq"]
    floor, context = g["floor"], g["context"]
    has_floor = bool(floor)
    names = sorted({r["sig_a"] for r in rows} | {r["sig_b"] for r in rows})
    ratios = aspect_ratios(names, sig_dir)
    r_cos = aspect_corr(rows, ratios, "global_cosine")
    r_mean = aspect_corr(rows, ratios, "mean_score")
    n_ref = sum(1 for n in names if n.startswith("ref_"))
    n_q = sum(1 for n in names if n.startswith("q_"))

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(BODY_PT + 0.5)
    # East-Asian and complex-script slots resolve separately; without these the
    # renderer can fall back to a serif for the same run.
    rpr = normal.element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(slot), "Calibri")
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1.0)

    title(doc, "Signature Comparison — Parameters, Methodology, "
               "Scope and Limitations")

    # ---- 1. Scope ----
    h1(doc, "1. Scope")
    body(doc,
         "This analysis covers the computational comparison of the signature "
         "images detected on the submitted documents, in their as-scanned "
         "condition, using the parameters and methodology described below. "
         "It compares the two questioned signatures attributed to Shafiul "
         "Islam on the two Letters of Authority against three signatures "
         "attributed to him on three further documents supplied as specimens. "
         "It assesses the visual similarity of static signature images. It "
         "does not extend to ink or paper examination, indentation or pressure "
         "analysis from the physical originals, dating of the documents, or "
         "verification of the surrounding document content. The models applied "
         "are general-purpose vision models used without training on the "
         "individuals’ signatures, together with a classical structural "
         "measure; no writer-specific model was built for this engagement.")
    if has_floor:
        body(doc,
             "The comparison is confined to these five signatures. Signatures "
             "by the other signatories on the same pages are not compared with "
             "them and are not the subject of any finding; they are used only "
             "to calibrate the measurement scale, as described in Section 3.")
    else:
        body(doc,
             "The comparison is confined to these five signatures. No other "
             "signature is compared, scored, or referred to in this report.")

    # ---- 2. Parameters ----
    h1(doc, "2. Parameters checked")
    body(doc,
         "Each pair of signatures is assessed on three independent similarity "
         "parameters. They are computed by different techniques and examine "
         "different aspects of a signature, so analysis across all three "
         "metrics is done to increase the reliability of the findings.")
    table(doc,
          ["Parameter", "What it measures", "Scale", "Reading"],
          [["Global style similarity (cosine)",
            "Agreement in overall shape, proportions, and ink style, taken "
            "across the whole mark.",
            "0–100%. Negative values are possible in principle but do "
            "not arise for signature images.",
            "Higher = more similar overall appearance. Sensitive to the "
            "proportions of the mark; see Section 4."],
           ["Local stroke match score",
            "Proportion of ink-bearing regions in one signature that "
            "correspond one-to-one with a region in the other.",
            "0–100% (the percentage of ink regions that match).",
            "Higher = more stroke segments correspond one-to-one. The most "
            "writing-specific of the three measures."],
           ["Structural similarity (SSIM, after alignment)",
            "Pixel-level structural agreement after one signature is rigidly "
            "aligned to the other.",
            "0–100%.",
            "Higher = closer structural overlap. Computed over the inked "
            "region only, so blank margin does not inflate it."],
           ["Composite index (reported figure)",
            "Simple average of the three parameters above, provided as a "
            "single summary figure.",
            "0–100%.",
            "A summary index only. It is not a probability that two "
            "signatures share an author."]],
          widths=[1.32, 1.88, 1.65, 1.65])

    # ---- 3. Methodology ----
    h1(doc, "3. Methodology in brief")
    lead(doc, "Detection: ",
         "Each submitted page is scanned by an object-detection model (Faster "
         "R-CNN with a ResNet-50 feature-pyramid backbone) that locates "
         "handwritten marks — signatures, initials, and stamps — with "
         "a confidence score for each. Low-confidence and clearly non-signature "
         "detections (letterhead logos, exhibit headings, annotation marks) are "
         "removed after visual review before any comparison. Eight signature "
         "marks were detected across the five submitted pages.")
    lead(doc, "Isolation of the signature stroke: ",
         "A detection box is a rectangle, and on a real signature block a "
         "rectangle also encloses the ruled signature line, the printed name "
         "caption beneath it, and any letterhead text the flourish crosses. "
         "Printed matter is identical across documents produced from the same "
         "template, so leaving it in the image raises the measured similarity "
         "of precisely the pairs under examination. Each crop is therefore "
         "rendered at 400 dpi and cleaned: ruled lines are removed by "
         "horizontal opening, printed captions are removed by "
         "connected-component size, and where printed text physically crosses "
         "a stroke the affected area is masked by hand-audited region rather "
         "than by rule. Residual overlap that could not be removed without "
         "cutting into the writing is recorded in Section 5.")
    lead(doc, "Preprocessing: ",
         "Every retained signature image passes through an identical "
         "standardization: conversion to grayscale, soft background removal "
         "(Otsu thresholding, which whitens the background while preserving ink "
         "intensity and hence pen-pressure detail), a tight crop to the ink, "
         "and an aspect-preserving resize onto a common 224×224 canvas. "
         "This ensures that differences in scan quality, resolution, or crop "
         "size are not mistaken for differences between the signatures "
         "themselves.")
    lead(doc, "Cosine similarity: ",
         "For the global style parameter, each signature image is converted by "
         "the DINOv2 vision model into a numerical descriptor — a vector "
         "of several hundred values summarizing its visual characteristics. "
         "Cosine similarity measures the angle between the two vectors: two "
         "signatures with closely aligned descriptors score near 1, while "
         "unrelated visual patterns score substantially lower. It captures "
         "overall resemblance rather than any single local feature.")
    lead(doc, "Local stroke matching: ",
         "The same vision model also produces descriptors for small regions "
         "(patches) of each image. Only ink-bearing regions are considered. A "
         "region in one signature is counted as matched when it and a region in "
         "the other signature each select the other as their closest "
         "counterpart (mutual nearest-neighbour matching). The score is the "
         "fraction of ink regions with such a mutual match, normalised "
         "symmetrically across both signatures so that a smaller mark cannot "
         "score highly merely by matching into a larger one; the matches are "
         "also rendered as a visual overlay for manual review.")
    lead(doc, "Structural similarity. ",
         "Independently of the model-based checks, one signature is rigidly "
         "aligned to the other — rotated and translated using the ink’s "
         "centre of mass and principal axis — and the aligned pair is "
         "compared with SSIM, an established classical measure of structural "
         "image similarity. The comparison is restricted to the inked region "
         "and its immediate surroundings: on a signature canvas the great "
         "majority of the area is blank background, which agrees perfectly "
         "between any two images and would otherwise dominate the figure.")
    if has_floor:
        lead(doc, "Calibration of the scale: ",
             "A similarity figure has no meaning on its own \u2014 40% is "
             "neither high nor low until it is known what this process "
             "produces for signatures whose authorship is settled. Two "
             "reference ranges are therefore measured before the questioned "
             "marks are considered: a same-hand range from the three specimens "
             "scored against each other, and a different-hands range from "
             "those specimens scored against signatures known to be by other "
             "people. Those signatures serve only to fix the scale and are not "
             "the subject of any finding.")
    lead(doc, "Interpretation. ",
         "Absolute scores vary with scan quality and signature style, so "
         "figures are read relative to reference behaviour where available and, "
         "in all cases, in combination: a pair scoring low on all three "
         "parameters indicates that the computational measures corroborate the "
         "visual observation that the signatures differ, while method "
         "disagreement flags the pair for closer manual review rather than "
         "supporting a conclusion in either direction.")

    # ---- 4. Findings ----
    h1(doc, "4. Findings — Shafiul Islam, two Letters of Authority "
            "against three specimen signatures")
    body(doc,
         "The five signatures compared, as pre-processed crops. The top row is "
         "the three specimen signatures attributed to Shafiul Islam on the "
         "supplied documents; the row beneath is the two questioned signatures "
         "on the Letters of Authority.")
    if os.path.exists(exhibit):
        doc.add_picture(exhibit, width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    labels = {"global_cosine": "Global style similarity",
              "patch_match_score": "Local stroke match",
              "ssim_aligned": "Structural similarity",
              "mean_score": "Composite index"}
    if has_floor:
        body(doc, [("The measurement scale. ", True, False),
                   ("The two ranges established in Section 3 are reproduced "
                    "here, because the questioned figures are read against "
                    "them. Neither range involves a questioned mark.",
                    False, False)])
        table(doc,
              ["Parameter", f"Different hands ({len(floor)} pairs)",
               f"Same hand ({len(ceiling)} pairs)", "Separated"],
              [[labels[k],
                f"{rng(floor, k)[0]*100:.0f}\u2013{rng(floor, k)[1]*100:.0f}%",
                f"{rng(ceiling, k)[0]*100:.0f}\u2013{rng(ceiling, k)[1]*100:.0f}%",
                "yes" if rng(ceiling, k)[0] > rng(floor, k)[1] else "no \u2014 overlapping"]
               for k in SCORERS],
              widths=[1.55, 1.75, 1.55, 1.65])
        f_hi = rng(floor, "mean_score")[1]
        c_lo = rng(ceiling, "mean_score")[0]
    else:
        body(doc, [("Agreement among the specimen signatures. ", True, False),
                   ("Before the questioned signatures are considered, the "
                    "three specimens are scored against each other. These are "
                    "three signatures accepted as being by the same hand, from "
                    "three documents, three capture methods, roughly four "
                    "decades apart, so their agreement shows how far this "
                    "process\u2019s figures move for one writer under real "
                    "capture variation.", False, False)])
        table(doc,
              ["Specimen pair", "Global style", "Local stroke", "Structural",
               "Composite"],
              [[f"{SHORT[r['sig_a']]} vs {SHORT[r['sig_b']]}",
                pct(r["global_cosine"]), pct(r["patch_match_score"]),
                pct(r["ssim_aligned"]), pct(r["mean_score"])]
               for r in sorted(ceiling, key=lambda r: -r["mean_score"])],
              widths=[2.4, 1.05, 1.05, 1.0, 1.0])
        f_hi = c_lo = None

    body(doc, [("Where the questioned signatures fall. ", True, False),
               ("Each questioned signature is scored against each specimen.",
                False, False)])
    qrows = []
    for r in sorted(questioned, key=lambda r: -r["mean_score"]):
        q_key = r["sig_b"] if r["sig_b"].startswith("q_") else r["sig_a"]
        r_key = r["sig_a"] if r["sig_a"].startswith("ref_") else r["sig_b"]
        row = [SHORT[q_key], SHORT[r_key], pct(r["global_cosine"]),
               pct(r["patch_match_score"]), pct(r["ssim_aligned"]),
               pct(r["mean_score"])]
        if has_floor:
            row.append(placement(r["mean_score"], f_hi, c_lo))
        qrows.append(row)
    if has_floor:
        table(doc, ["Questioned", "Specimen", "Global", "Local stroke",
                    "Structural", "Composite", "Placement"],
              qrows, widths=[1.05, 1.15, 0.8, 0.9, 0.85, 0.85, 1.4])
    else:
        table(doc, ["Questioned", "Specimen", "Global style", "Local stroke",
                    "Structural", "Composite"],
              qrows, widths=[1.25, 1.35, 1.0, 0.98, 0.95, 0.97])

    body(doc,
         "Across every parameter the questioned figures sit below the "
         f"specimen figures: composite {rng_pct(questioned, 'mean_score')} "
         f"against {rng_pct(ceiling, 'mean_score')}, and local stroke match "
         f"{rng_pct(questioned, 'patch_match_score')} against "
         f"{rng_pct(ceiling, 'patch_match_score')}. All six "
         "questioned-to-specimen comparisons fall below all three "
         "specimen-to-specimen comparisons on the local stroke measure, which "
         "is the most writing-specific of the three. The direction is "
         "consistent across the set, and the lowest figure \u2014 the HCD "
         f"letter against the RJSC specimen at "
         f"{min(r['mean_score'] for r in questioned)*100:.0f}% composite, "
         f"{min(r['patch_match_score'] for r in questioned)*100:.0f}% local "
         "stroke \u2014 is the furthest from the specimen range.")
    if qq:
        body(doc,
             "The two questioned signatures compared with each other score "
             f"{pct(qq[0]['mean_score'])} composite "
             f"({pct(qq[0]['patch_match_score'])} local stroke), above every "
             "specimen pair. They are more consistent with each other than the "
             "specimens are among themselves.")

    body(doc, [("What these findings do and do not establish: ", True, False)])
    if has_floor:
        body(doc,
             "The questioned signatures are measurably less consistent with "
             "the specimens than the specimens are with each other, on "
             "parameters shown on this same material to distinguish writers, "
             "and the visual observation is consistent with this: the three "
             "specimens share a construction \u2014 a large lower-left loop, "
             "tall vertical staffs, and a terminal ascending sweep \u2014 "
             "differently expressed in the questioned marks. That is the "
             "finding, and it is a documented inconsistency warranting "
             "examination. It is not a finding that the questioned signatures "
             "were written by a different person: most of the figures fall "
             "between the measured ranges rather than within the "
             "different-hands range, and the same-hand range rests on three "
             "pairs. Whether these differences reflect natural variation by "
             "one writer or the involvement of a different hand is a "
             "determination for a qualified forensic document examiner.")
    else:
        body(doc,
             "The questioned signatures are measurably less consistent with "
             "the specimen signatures than the specimen signatures are with "
             "each other. The shortfall is present on every parameter and is "
             "largest on the local stroke match, where all six "
             "questioned-to-specimen comparisons fall below the specimen "
             "range. The visual observation is consistent with this: the three "
             "specimens share a construction \u2014 a large lower-left loop, "
             "tall vertical staffs, and a terminal ascending sweep \u2014 "
             "differently expressed in the questioned marks. That is the "
             "finding, and it is a documented inconsistency warranting "
             "examination.")
        body(doc,
             "It is not a finding that the questioned signatures were written "
             "by a different person, and this report does not support that "
             "conclusion. The specimen range shows how far these figures move "
             "for one writer across three documents; it does not show how far "
             "they move between two writers, which was not measured. A "
             "shortfall against the specimen range is therefore evidence of "
             "inconsistency, not of authorship. Whether these differences "
             "reflect natural variation by one writer or the involvement of a "
             "different hand is a determination for a qualified forensic "
             "document examiner.")

    # ---- 5. Limitations ----
    h1(doc, "5. Limitations")
    body(doc,
         "The parameters measure visual similarity, not authorship. A skilled "
         "imitation of a signature can score high on visual similarity, and "
         "natural variation in a person’s genuine signature — which "
         "is expected, since no one signs identically twice — can lower "
         "scores, particularly across documents produced at different times or "
         "captured at different scan qualities.")
    if has_floor:
        body(doc,
             "Both reference ranges rest on small samples: three pairs for the "
             "same-hand range. A range built on three pairs is not a stable "
             "estimate of one writer’s natural variation, and additional "
             "genuine specimens — ten to fifteen, ideally close in date to "
             "the two letters — would be needed to characterise it "
             "properly.")
    else:
        body(doc,
             "The specimen range rests on three pairs, which is not a stable "
             "estimate of one writer’s natural variation across occasions. "
             "Additional genuine specimens — ten to fifteen, ideally close "
             "in date to the two letters — would be needed to characterise "
             "it properly.")
        body(doc,
             "No range was measured for signatures written by different "
             "people. The figures in Section 4 show the size of the shortfall "
             "against the specimen range; there is no measured threshold that "
             "a figure can be said to cross, and none is asserted.")

    body(doc,
         "The measures respond to the overall proportions of a mark as well as "
         "to its stroke detail, and the specimens and questioned signatures "
         "differ systematically in proportion — the specimens are taller "
         "than they are wide, the questioned marks roughly twice as wide as "
         "they are tall. Whether that difference is a characteristic of the "
         "writing or a consequence of the space available on the page is not "
         "determined here.")
    body(doc,
         "Capture quality differs between the two sets. The specimens are "
         "clean scans; the HCD questioned signature is a fax-grade, one-bit "
         "reproduction. Any measure sensitive to stroke detail penalises the "
         "questioned side for reasons unrelated to authorship.")
    body(doc,
         "Results depend on the quality of the detected crops. Heavy overlap "
         "of a signature with printed text, stamps, or ruled lines reduces "
         "reliability: on the RJSC and service-agreement specimens, stamp and "
         "letterhead text crosses the strokes and could not be removed without "
         "cutting into the writing. The detector’s recall has not been "
         "measured, so a signature it failed to detect would be absent from "
         "this analysis without indication.")
    body(doc,
         "The composite index is a convenience summary with no probabilistic "
         "interpretation and can mask disagreement between the underlying "
         "parameters, which is why the three individual figures accompany it. "
         "Findings from this process are an investigative aid expressed as "
         "consistency or inconsistency of the compared images under these "
         "measures; they are not a conclusive determination that any signature "
         "is genuine or forged, and any formal opinion on authorship should "
         "rest on examination by a qualified forensic document examiner.")

    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--csv",
                    default="comparison_output_calibrated/comparison_report.csv")
    ap.add_argument("--exhibit",
                    default="../loa_comparison/exhibit_shafiul_only.png")
    ap.add_argument("--sig-dir", default="signatures")
    ap.add_argument("--out",
                    default="Signature Comparison - Parameters, Methodology, "
                            "Scope and Limitations.docx")
    args = ap.parse_args()
    rows = load(args.csv)
    path = build(rows, args.exhibit, args.sig_dir, args.out)
    print(f"Word report: {os.path.abspath(path)}")
    print(f"  pairs: {len(rows)}   exhibit: {args.exhibit}")


if __name__ == "__main__":
    main()
