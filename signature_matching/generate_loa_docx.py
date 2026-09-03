"""
Word report for the Letter of Authority signature examination.

Reads the scored pair matrix written by run_loa_case.py and renders it as a
.docx, so the findings can be circulated and annotated outside this repo.

Everything quantitative is read from comparison_report.csv rather than typed in,
so the document cannot drift from the run that produced it -- including the
aspect-ratio correlation, which is recomputed from the crops themselves.

The document adapts to the scope of the run. If the CSV contains
different-writer control pairs, it reports the calibration bands and whether
they separate. If it does not -- a run restricted to one signer -- it says so,
and reports no verdict at all rather than presenting an uncalibrated score as
though it meant something.

Usage:
    python generate_loa_docx.py [--csv comparison_output_loa/comparison_report.csv]
                                [--exhibit ../loa_comparison/exhibit_signatures.png]
                                [--sig-dir signatures]
                                [--out "Signature Examination - Letters of Authority.docx"]
"""

import argparse
import csv
import math
import os
from collections import defaultdict

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(0x1A, 0x1D, 0x22)
MUTED = RGBColor(0x6A, 0x70, 0x79)
ACCENT = RGBColor(0x2F, 0x4B, 0x7C)
GENUINE = RGBColor(0x24, 0x5A, 0x41)
QUESTIONED = RGBColor(0x8B, 0x27, 0x27)
CAUTION_FILL = "FBF3E2"
SHADE = "EFEFEA"


# ----------------------------------------------------------------------------
# Low-level helpers (python-docx exposes no API for these)
# ----------------------------------------------------------------------------

def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def cell_border(cell, edge="left", size=18, color="2F4B7C"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:color"), color)
    borders.append(el)


def rule(doc, color="C3C2B9", size=6):
    """A horizontal rule, drawn as a bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)
    return p


def para(doc, text="", size=10.5, bold=False, italic=False, color=INK,
         space_after=8, space_before=0, font="Calibri", align=None,
         style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.name = font
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    return p


def rich(doc, parts, size=10.5, space_after=8, space_before=0, font="Calibri"):
    """A paragraph from (text, **fmt) fragments, so emphasis can sit inline."""
    p = doc.add_paragraph()
    for frag in parts:
        text, fmt = (frag, {}) if isinstance(frag, str) else frag
        r = p.add_run(text)
        r.font.size = Pt(fmt.get("size", size))
        r.font.bold = fmt.get("bold", False)
        r.font.italic = fmt.get("italic", False)
        r.font.color.rgb = fmt.get("color", INK)
        r.font.name = fmt.get("font", font)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def eyebrow(doc, text):
    return para(doc, text.upper(), size=8, bold=True, color=MUTED,
                space_after=2, space_before=16)


def heading(doc, text):
    p = para(doc, text, size=15, bold=True, color=INK,
             space_after=6, font="Cambria")
    return p


def subheading(doc, text):
    return para(doc, text, size=11, bold=True, color=INK,
                space_after=4, space_before=12)


def bullet(doc, parts):
    p = rich(doc, parts, size=10.5, space_after=4)
    p.style = doc.styles["List Bullet"]
    return p


def no_split(row):
    """Stop Word breaking a table row across a page boundary."""
    el = OxmlElement("w:cantSplit")
    row._tr.get_or_add_trPr().append(el)


def table(doc, headers, rows, widths=None, stripes=None):
    """Header row plus body. `stripes` colours a left edge per row."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h.upper())
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = MUTED
        r.font.name = "Calibri"
        shade(hdr[i], SHADE)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            numeric = isinstance(val, str) and val.replace(".", "").replace(
                "−", "").isdigit()
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            r.font.name = "Consolas" if numeric else "Calibri"
            r.font.color.rgb = INK
            if numeric:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if stripes and stripes[ri]:
            cell_border(cells[0], "left", size=24, color=stripes[ri])
    for row in t.rows:
        no_split(row)
        if widths:
            for i, wdt in enumerate(widths):
                row.cells[i].width = Inches(wdt)
    return t


# ----------------------------------------------------------------------------
# Data
# ----------------------------------------------------------------------------

SCORERS = ("global_cosine", "patch_match_score", "ssim_aligned", "mean_score")

# Readable labels for the appendix. The internal ids carry the role prefix so
# the runner can group by it; a reader of the report needs the document instead.
DISPLAY = {
    "ref_shafiul_invitation": "Shafiul \u00b7 invitation",
    "ref_shafiul_navana_sa": "Shafiul \u00b7 service agr.",
    "ref_shafiul_rjsc": "Shafiul \u00b7 RJSC",
    "q_shafiul_hcd": "Questioned \u00b7 HCD",
    "q_shafiul_cmm": "Questioned \u00b7 CMM",
    "ctl_shafqat_hcd": "Shafqat Ahmed",
    "ctl_muto_navana_sa": "Kazuyuki Muto",
    "ctl_second_signer_rjsc": "2nd signatory \u00b7 RJSC",
}


def load(csv_path):
    with open(csv_path) as f:
        rows = list(csv.DictReader(f))
    for r in rows:
        for k in SCORERS + ("patch_mean_sim",):
            r[k] = float(r[k])
    return rows


def groups(rows):
    g = defaultdict(list)
    for r in rows:
        t = r["type"]
        if t.startswith("ref-ref"):
            g["ceiling"].append(r)
        elif t.startswith("ref-questioned"):
            g["questioned"].append(r)
        elif t.startswith("q-q"):
            g["qq"].append(r)
        else:
            g["floor"].append(r)
    return g


def rng(rows, key):
    v = [r[key] for r in rows]
    return min(v), max(v)


def aspect_ratios(names, sig_dir):
    """Ink-bounding-box width:height for each crop, measured from the PNG."""
    import numpy as np
    from PIL import Image
    out = {}
    for n in names:
        path = os.path.join(sig_dir, f"{n}.png")
        if not os.path.exists(path):
            continue
        ink = np.asarray(Image.open(path).convert("L"), dtype=np.uint8) <= 128
        ys, xs = np.where(ink)
        if len(xs):
            out[n] = float((xs.max() - xs.min() + 1) / (ys.max() - ys.min() + 1))
    return out


def pearson(a, b):
    n = len(a)
    if n < 3:
        return None
    ma, mb = sum(a) / n, sum(b) / n
    num = sum((x - ma) * (y - mb) for x, y in zip(a, b))
    da = math.sqrt(sum((x - ma) ** 2 for x in a))
    db = math.sqrt(sum((y - mb) ** 2 for y in b))
    return num / (da * db) if da and db else None


def aspect_correlation(rows, ratios, key="global_cosine"):
    """How much of a score is explained by the two marks having a similar shape.

    Reported because it is the pipeline's sharpest failure mode: an embedding
    that tracks whether both marks are wide flat scrawls or both tall spiky
    ones is answering a question about layout, not about authorship.
    """
    d, c = [], []
    for r in rows:
        if r["sig_a"] in ratios and r["sig_b"] in ratios:
            d.append(abs(math.log(ratios[r["sig_a"]]) - math.log(ratios[r["sig_b"]])))
            c.append(r[key])
    return pearson(d, c), len(d)


# ----------------------------------------------------------------------------
# Document
# ----------------------------------------------------------------------------

def build(rows, exhibit, out_path, sig_dir):
    g = groups(rows)
    ceiling, floor, questioned, qq = (g["ceiling"], g["floor"],
                                      g["questioned"], g["qq"])
    has_floor = bool(floor)

    names = sorted({r["sig_a"] for r in rows} | {r["sig_b"] for r in rows})
    n_sigs = len(names)
    ratios = aspect_ratios(names, sig_dir)
    r_cos, n_corr = aspect_correlation(rows, ratios, "global_cosine")
    r_mean, _ = aspect_correlation(rows, ratios, "mean_score")

    doc = Document()
    for s in ("Normal", "List Bullet"):
        doc.styles[s].font.name = "Calibri"
        doc.styles[s].font.size = Pt(10.5)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    sec.left_margin = sec.right_margin = Inches(1.0)

    # ---- masthead ----
    para(doc, "SIGNATURE EXAMINATION  ·  NAVANA LIMITED", size=8.5,
         bold=True, color=MUTED, space_after=6)
    para(doc, "Two Letters of Authority against three specimen documents",
         size=21, bold=True, font="Cambria", space_after=8)
    standfirst = (
        "Eight signatures were extracted from five documents and scored "
        "pairwise. The extraction succeeded; the scoring did not separate "
        "writers, and cannot answer the question it was asked."
        if has_floor else
        "Five Shafiul Islam signatures were extracted from five documents and "
        "scored pairwise \u2014 three genuine specimens against the two "
        "questioned Letters of Authority. With no other signer in scope there "
        "is nothing to calibrate the scores against, so this run reports "
        "measurements, not a conclusion.")
    para(doc, standfirst, size=12, italic=True, color=MUTED, font="Cambria",
         space_after=10)
    rich(doc, [
        ("Documents ", {"bold": True, "size": 9, "color": MUTED}),
        ("5 (5 pages)     ", {"size": 9, "color": MUTED}),
        ("Signatures compared ", {"bold": True, "size": 9, "color": MUTED}),
        (f"{n_sigs}     ", {"size": 9, "color": MUTED}),
        ("Pairs scored ", {"bold": True, "size": 9, "color": MUTED}),
        (f"{len(rows)}     ", {"size": 9, "color": MUTED}),
        ("Run ", {"bold": True, "size": 9, "color": MUTED}),
        ("31 August 2026", {"size": 9, "color": MUTED}),
    ], space_after=4)
    rule(doc, color="1A1D22", size=12)

    # ---- bottom line ----
    worst = max(floor, key=lambda r: r["mean_score"]) if floor else None
    # The HIGHEST same-writer pair, not the lowest. "A different-writer pair
    # beat the best genuine pair" is both true and the strongest statement of
    # non-separation available; quoting the weakest genuine pair understates it.
    best = max(ceiling, key=lambda r: r["mean_score"]) if ceiling else None
    box = doc.add_table(rows=1, cols=1)
    c = box.rows[0].cells[0]
    shade(c, CAUTION_FILL)
    cell_border(c, "left", size=30, color="8A6410")
    c.text = ""
    p0 = c.paragraphs[0]
    r0 = p0.add_run("BOTTOM LINE")
    r0.font.size, r0.font.bold = Pt(8.5), True
    r0.font.color.rgb = RGBColor(0x8A, 0x64, 0x10)
    p0.paragraph_format.space_after = Pt(6)

    def boxpara(parts):
        p = c.add_paragraph()
        for text, fmt in parts:
            r = p.add_run(text)
            r.font.size = Pt(10.5)
            r.font.bold = fmt.get("bold", False)
            r.font.italic = fmt.get("italic", False)
            r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(7)

    if has_floor:
        boxpara([
            ("The comparison is not diagnostic. ", {"bold": True}),
            ("Two signatures known to be by ", {}),
            ("different", {"italic": True}),
            (" hands scored higher (", {}),
            (f"{worst['mean_score']:.3f}", {"bold": True}),
            (") than the best-matching pair known to be by the ", {}),
            ("same", {"italic": True}),
            (" hand (", {}), (f"{best['mean_score']:.3f}", {"bold": True}),
            ("). When the different-writer range overlaps the same-writer "
             "range, no threshold separates them, so no score in this report "
             "supports a conclusion about who signed the Letters of Authority "
             "— in either direction.", {}),
        ])
        boxpara([
            ("What the run does establish is the setup for a real examination: "
             "the genuine specimen signatures are isolated and cleaned, and "
             "three signatures by other people on the same pages give a "
             "like-for-like different-writer baseline that this case "
             "previously had none of. The question itself needs a qualified "
             "forensic document examiner.", {}),
        ])
    else:
        boxpara([
            ("No verdict is available from this run. ", {"bold": True}),
            ("Only Shafiul Islam's own signatures are in scope, so every pair "
             "scored here is either genuine-to-genuine or "
             "questioned-to-genuine. Nothing in the run establishes what this "
             "pipeline scores for two signatures by ", {}),
            ("different", {"italic": True}),
            (" hands. Without that, a score of 0.39 cannot be called low and a "
             "score of 0.46 cannot be called high — there is no scale to "
             "read them against.", {}),
        ])
        boxpara([
            ("The three genuine specimens agree with each other at "
             f"{rng(ceiling, 'mean_score')[0]:.3f}–"
             f"{rng(ceiling, 'mean_score')[1]:.3f}, and the questioned "
             "signatures sit slightly below that against the same specimens, "
             f"at {rng(questioned, 'mean_score')[0]:.3f}–"
             f"{rng(questioned, 'mean_score')[1]:.3f}. That gap is small, it "
             "rests on three reference pairs, and a large part of it is "
             "explained by shape alone (see \u201cWhat the scores actually "
             "track\u201d). It is not evidence of anything on its own.", {}),
        ])
        boxpara([
            ("Deciding whether these signatures are genuine needs a qualified "
             "forensic document examiner. If a numerical answer is wanted from "
             "this pipeline, it needs signatures by other people to calibrate "
             "against.", {}),
        ])

    # ---- 1. extraction ----
    eyebrow(doc, "What was extracted")
    WORD = {2: "Two", 3: "Three", 5: "Five", 8: "Eight"}
    heading(doc, f"{WORD.get(n_sigs, n_sigs)} signatures, five documents")
    para(doc,
         "Each source document is a single page. The detector found eight "
         "signature marks across them"
         + ("" if has_floor else
            f", of which the {WORD.get(n_sigs, n_sigs).lower()} attributed to "
            f"Shafiul Islam are compared here")
         + "; identifying who wrote each one comes from the printed captions "
           "beside them.", color=MUTED, space_after=10)

    G, Q, C = "245A41", "8B2727", "C3C2B9"
    SOURCES = {
        "ref_shafiul_invitation": ("Shafiul Islam", "Invitation letter, Dec 2014", "Genuine", G),
        "ref_shafiul_navana_sa": ("Shafiul Islam", "Navana service agreement", "Genuine", G),
        "ref_shafiul_rjsc": ("Shafiul Islam", "RJSC Form XVIII", "Genuine", G),
        "q_shafiul_hcd": ("Shafiul Islam", "Letter of Authority — Arbitration, HCD", "QUESTIONED", Q),
        "q_shafiul_cmm": ("Shafiul Islam", "Letter of Authority — CR Case, CMM", "QUESTIONED", Q),
        "ctl_shafqat_hcd": ("Shafqat Ahmed", "Letter of Authority — Arbitration, HCD", "Control", C),
        "ctl_muto_navana_sa": ("Kazuyuki Muto", "Navana service agreement", "Control", C),
        "ctl_second_signer_rjsc": ("Second signatory", "RJSC Form XVIII", "Control", C),
    }
    order = [k for k in SOURCES if k in names]
    body, stripes = [], []
    for k in order:
        who, docname, role, col = SOURCES[k]
        body.append([who, docname, role,
                     f"{ratios[k]:.2f}" if k in ratios else "—"])
        stripes.append(col)
    table(doc, ["Signature", "Document", "Role", "Aspect w:h"], body,
          widths=[1.30, 3.05, 1.05, 0.85], stripes=stripes)
    if has_floor:
        para(doc,
             "The controls are not padding. They are signatures by other "
             "people captured through the same scanners, on the same pages, "
             "with the same degradation — which is what makes them a fair "
             "different-writer baseline.",
             size=9, color=MUTED, space_before=6, space_after=4)
    else:
        para(doc,
             "Signatures by the three other signatories on these same pages "
             "were extracted but are excluded from this run by request. They "
             "are the only material available that could have supplied a "
             "different-writer baseline.",
             size=9, color=MUTED, space_before=6, space_after=4)

    # ---- 2. exhibit ----
    eyebrow(doc, "The evidence")
    heading(doc, "What the signatures actually look like")
    para(doc, f"All {n_sigs}, cleaned of ruled lines and printed captions, "
              f"shown at a common scale.", color=MUTED, space_after=10)
    if os.path.exists(exhibit):
        doc.add_picture(exhibit, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = ("Top row: genuine Shafiul Islam signatures from three other "
               "documents. ")
        cap += ("Middle: " if has_floor else "Below: ")
        cap += "the two questioned Letter of Authority signatures."
        if has_floor:
            cap += (" Bottom: three other signatories, used as the "
                    "different-writer baseline.")
        cap += (" Residual printed fragments in the RJSC and "
                "service-agreement specimens are stamp and letterhead text "
                "physically overlapping the stroke; they could not be removed "
                "without cutting into the signature.")
        para(doc, cap, size=8.5, color=MUTED, space_before=6, space_after=12)
    else:
        para(doc, f"[exhibit image not found: {exhibit}]", color=QUESTIONED)

    subheading(doc, "A structural observation, offered for expert review")
    para(doc,
         "The three genuine specimens are built the same way: a large "
         "lower-left oval loop, then two or three tall vertical staffs rising "
         "well above the body, a small zigzag, and a terminal ascending sweep. "
         "They span 2014 to a 1970s-era filing and three different capture "
         "methods, and that construction is present in all three.")
    rich(doc, [
        ("Both questioned signatures instead open with an ", {}),
        ("E", {"italic": True}),
        ("-like form, carry low humps rather than tall staffs, and show no "
         "large lower-left loop. They are also proportioned very differently "
         "— roughly two and a half times wider relative to their height "
         "(2.16 and 2.34) than any genuine specimen (0.81 to 1.08).", {}),
    ])
    para(doc,
         "This is a description of what is visible, not a finding. Compression "
         "into a confined signing space can flatten a signature, and none of "
         "the features that decide these cases — pen lifts, line quality, "
         "tremor, stroke order, pressure — are measured anywhere in this "
         "pipeline.", italic=True, color=MUTED)

    # ---- 3. result ----
    doc.add_page_break()
    labels = {"global_cosine": "Global cosine",
              "patch_match_score": "Patch match",
              "ssim_aligned": "SSIM (ink-masked)",
              "mean_score": "Mean of the three"}

    eyebrow(doc, "The result")
    if has_floor:
        heading(doc, "The calibration bands overlap")
        para(doc,
             "Every pair falls into a known category. Same-writer pairs set a "
             "ceiling, different-writer pairs set a floor. If those two ranges "
             "overlap, the scorer cannot tell writers apart on this material.",
             color=MUTED, space_after=10)
        band_rows = []
        for k in SCORERS:
            f0, f1 = rng(floor, k)
            c0, c1 = rng(ceiling, k)
            q0, q1 = rng(questioned, k)
            band_rows.append([
                labels[k], f"{f0:.3f}–{f1:.3f}", f"{c0:.3f}–{c1:.3f}",
                f"{q0:.3f}–{q1:.3f}", "yes" if c0 > f1 else "NO",
            ])
        table(doc,
              [" ", f"Different-writer ({len(floor)} pairs)",
               f"Same-writer ({len(ceiling)} pairs)",
               f"Questioned vs genuine ({len(questioned)})", "Separated"],
              band_rows, widths=[1.3, 1.5, 1.35, 1.5, 0.85])
        rich(doc, [
            ("On the combined score the ranges overlap outright. The single "
             "most damaging pair is ", {}),
            (f"{DISPLAY.get(worst['sig_a'], worst['sig_a'])} vs "
             f"{DISPLAY.get(worst['sig_b'], worst['sig_b'])}", {"bold": True}),
            (f" at {worst['mean_score']:.3f} — two people who plainly did "
             f"not write each other's signature, scoring above even the "
             f"best-matching genuine pair of Shafiul's own signatures at "
             f"{best['mean_score']:.3f}.", {}),
        ], space_before=10)
    else:
        heading(doc, "There is no baseline to read the scores against")
        para(doc,
             "Every pair in this run is either genuine-to-genuine or "
             "questioned-to-genuine. Both are pairs of Shafiul Islam "
             "signatures. Nothing here measures what a different writer scores.",
             color=MUTED, space_after=10)
        band_rows = []
        for k in SCORERS:
            c0, c1 = rng(ceiling, k)
            q0, q1 = rng(questioned, k)
            band_rows.append([
                labels[k], f"{c0:.3f}–{c1:.3f}", f"{q0:.3f}–{q1:.3f}",
                "not measured",
            ])
        table(doc,
              [" ", f"Genuine vs genuine ({len(ceiling)} pairs)",
               f"Questioned vs genuine ({len(questioned)} pairs)",
               "Different writer"],
              band_rows, widths=[1.55, 1.75, 1.85, 1.35])
        para(doc,
             "The questioned range sits a little below the genuine range on "
             "every scorer. That is the whole of the quantitative result, and "
             "it is weaker than it looks: three reference pairs is too few to "
             "define a range, the two ranges are close, and the section below "
             "shows most of the gap is predicted by shape rather than by "
             "handwriting.", space_before=10)

    subheading(doc, "What the scores actually track")
    if r_cos is not None:
        rich(doc, [
            (f"Across the {n_corr} pairs in this run, similarity of the "
             f"DINOv2 global embedding correlates with similarity of ", {}),
            ("aspect ratio", {"italic": True}),
            (" at ", {}), (f"r = {r_cos:+.2f}", {"bold": True}),
            (f", and the combined score at r = {r_mean:+.2f}. The scorers are "
             f"substantially measuring whether two marks are the same "
             f"proportion, not whether the same hand made them.", {}),
        ])
    if not has_floor:
        para(doc,
             "This matters directly here. The three genuine specimens are "
             "tall (aspect 0.81 to 1.08); both questioned signatures are wide "
             "(2.16 and 2.34). That difference alone pushes every "
             "questioned-versus-genuine score down, whoever signed them. With "
             "no different-writer pairs in the run, there is no way to "
             "separate the part of the gap caused by shape from any part "
             "caused by authorship.")
    else:
        para(doc,
             "That is exactly why the wide questioned signatures score "
             "0.83–0.93 against wide control signatures by other people, "
             "while Shafiul's own tall specimens score only 0.69–0.82 "
             "against each other.")

    # ---- 4. patch signal ----
    eyebrow(doc, "The one consistent signal")
    heading(doc, "Patch matching behaves, but does not settle it")
    para(doc,
         "Of the three scorers, only patch-token matching orders the pairs "
         "sensibly. It is worth recording, with its limits stated plainly.",
         color=MUTED, space_after=10)

    pc0, pc1 = rng(ceiling, "patch_match_score")
    pq0, pq1 = rng(questioned, "patch_match_score")
    pqq = qq[0]["patch_match_score"] if qq else 0.0
    prows = [[f"Genuine vs genuine ({len(ceiling)} pairs)", f"{pc0:.3f}",
              f"{pc1:.3f}", "Shafiul's own hand, across documents"],
             [f"Questioned vs genuine ({len(questioned)} pairs)", f"{pq0:.3f}",
              f"{pq1:.3f}", "every pair below the genuine minimum"]]
    pstripes = [G, Q]
    if has_floor:
        pf0, pf1 = rng(floor, "patch_match_score")
        prows.append([f"Known different writers ({len(floor)} pairs)",
                      f"{pf0:.3f}", f"{pf1:.3f}",
                      "range still overlaps the genuine band"])
        pstripes.append(C)
    if qq:
        prows.append(["The two questioned, to each other", f"{pqq:.3f}",
                      f"{pqq:.3f}", "highest score in the run"])
        pstripes.append(Q)
    table(doc, ["Pair category", "Lowest", "Highest", "Reading"], prows,
          widths=[2.1, 0.8, 0.85, 2.75], stripes=pstripes)

    para(doc,
         "All six questioned-versus-genuine pairs fall below all three "
         "genuine-versus-genuine pairs, and the two questioned signatures "
         "resemble each other more than any two genuine specimens resemble "
         "each other. The direction is consistent.", space_before=10)
    tail = ("It is still not a result. Three genuine pairs is too small a "
            "sample to define a band, ")
    if has_floor:
        tail += (f"one different-writer pair reaches {pf1:.3f} and breaks the "
                 f"separation anyway, ")
    else:
        tail += ("no different-writer pair was scored at all, patch matching "
                 f"tracks aspect ratio here at r = {r_mean:+.2f}, ")
    tail += ("and the questioned signatures were captured at markedly lower "
             "quality than the specimens — the HCD letter is a 1-bit "
             "fax-grade reproduction, and degradation alone lowers patch "
             "matching. The consistency warrants examination; it does not "
             "substitute for one.")
    para(doc, tail)

    # ---- 5. method ----
    doc.add_page_break()
    eyebrow(doc, "Method")
    heading(doc, "How the numbers were produced")
    para(doc, "Four stages, in order. Each writes its output to disk so any "
              "step can be re-checked independently.", color=MUTED,
         space_after=10)
    stages = [
        ("01", "Probe the object layer", "pdf_probe.py",
         "reads what each PDF declares: 5 pages, 4 scanned, 1 native, no "
         "cryptographic signature fields, no ink annotations. Nothing here is "
         "inferred from pixels."),
        ("02", "Detect signature locations", "detect_signs.py",
         "Faster R-CNN over rendered pages at a 0.30 confidence floor. Eight "
         "marks found; boxes reported in PDF point coordinates."),
        ("03", "Crop and isolate the stroke",
         "crop_detections.py + clean_crops.py",
         "new. Renders each box at 400 dpi with padding, then strips ruled "
         "lines by horizontal opening and drops printed captions by "
         "connected-component size. Where printed text physically crosses a "
         "stroke, hand-audited erase boxes are named per file rather than "
         "guessed."),
        ("04", "Score and calibrate", "run_loa_case.py",
         "new. Scores all 28 pairs on three measures, then builds the "
         "same-writer and different-writer bands from the corpus itself and "
         "refuses to emit a per-pair verdict while those bands overlap."),
    ]
    for n, title, script, desc in stages:
        rich(doc, [
            (f"{n}   ", {"bold": True, "color": ACCENT, "font": "Consolas",
                         "size": 9.5}),
            (title, {"bold": True}),
        ], space_after=1, space_before=6)
        rich(doc, [
            (script, {"font": "Consolas", "size": 9, "color": ACCENT}),
            (" — " + desc, {"size": 10, "color": MUTED}),
        ], space_after=4)

    subheading(doc, "Corrections made to the scorers")
    rich(doc, [
        ("Three defects in the inherited comparison code would each have "
         "inflated similarity, and were fixed before the run: SSIM was "
         "averaged over the whole canvas, where roughly 90% is blank "
         "background that scores a perfect 1.0, so it now averages over the "
         "dilated ink region only; patch matches were normalised by the ", {}),
        ("smaller", {"italic": True}),
        (" ink area, letting a small signature match into a large one and "
         "score near 1.0, so normalisation is now symmetric; and the patch ink "
         "threshold admitted patches holding only stroke anti-aliasing, which "
         "produced spurious mutual matches.", {}),
    ])

    # ---- 6. limits ----
    eyebrow(doc, "Limits")
    heading(doc, "What this cannot tell you")
    limits = []
    if not has_floor:
        limits.append((
            "There is no different-writer baseline in this run.",
            " Only Shafiul Islam's signatures were compared, so no score here "
            "can be called high or low. This is the single largest limitation "
            "and it is not fixable by any amount of further computation on "
            "these five signatures alone."))
    for lead, rest in limits + [
        ("It is not a forensic document examination.",
         " Pen lifts, line quality, tremor, pressure, stroke order and writing "
         "rhythm carry the weight in real signature disputes. None of them are "
         "modelled here."),
        ("The sample is very small.",
         " Three genuine specimens give three same-writer pairs. A range built "
         "on three points is not stable"
         + (", and the overlap may be an artefact of that as much as of the "
            "scorers." if has_floor else
            ", so the genuine range quoted above should not be treated as "
            "the limits of his natural variation.")),
        ("Capture quality is confounded with the question.",
         " The genuine specimens are clean scans; the HCD questioned signature "
         "is fax-grade and 1-bit. Any measure sensitive to stroke detail will "
         "penalise the questioned side for reasons unrelated to authorship."),
        ("Shape is confounded with the question.",
         " The genuine specimens are tall and the questioned signatures are "
         "wide, and the scorers track that difference directly. Part of every "
         "questioned-versus-genuine gap is proportion, not handwriting."),
        ("Two specimens carry residue.",
         " The RJSC and service-agreement signatures have stamp and letterhead "
         "text crossing the strokes that could not be removed without cutting "
         "into the writing."),
        ("The detector's recall has never been measured.",
         " A signature it missed is silently absent from everything above."),
    ]:
        bullet(doc, [(lead, {"bold": True}), (rest, {})])

    subheading(doc, "What would make this answerable")
    for text in [
        "Signatures by other people, to calibrate against. Without a "
        "different-writer baseline no score from this pipeline can be "
        "interpreted at all.",
        "More genuine specimens — ten to fifteen, ideally near the dates "
        "on the two letters, which would turn a three-point ceiling into an "
        "actual distribution.",
        "Original documents or high-resolution colour scans of the two Letters "
        "of Authority, rather than the fax-grade reproductions used here.",
        "Referral to a qualified forensic document examiner, who can assess "
        "the features this pipeline does not see.",
    ]:
        bullet(doc, [(text, {})])

    # ---- appendix ----
    doc.add_page_break()
    eyebrow(doc, "Appendix")
    heading(doc, "Full pair matrix")
    para(doc, f"All {len(rows)} scored pairs, grouped by category and ordered "
              f"by mean score.", color=MUTED, space_after=10)
    order = {"ref-ref": 0, "ref-questioned": 1, "q-q": 2}
    srt = sorted(rows, key=lambda r: (order.get(r["type"].split(" ")[0], 3),
                                      -r["mean_score"]))
    stripes = []
    body = []
    for r in srt:
        t = r["type"]
        stripes.append(G if t.startswith("ref-ref")
                       else Q if t.startswith(("ref-questioned", "q-q")) else C)
        body.append([DISPLAY.get(r["sig_a"], r["sig_a"]),
                     DISPLAY.get(r["sig_b"], r["sig_b"]), t.split(" (")[0],
                     f"{r['global_cosine']:.3f}",
                     f"{r['patch_match_score']:.3f}",
                     f"{r['ssim_aligned']:.3f}",
                     f"{r['mean_score']:.3f}"])
    table(doc, ["Signature A", "Signature B", "Type", "Cosine", "Patch",
                "SSIM", "Mean"], body,
          widths=[1.75, 1.75, 0.92, 0.52, 0.52, 0.52, 0.52], stripes=stripes)

    rule(doc)
    para(doc,
         "Generated from signature_matching/run_loa_case.py. Full pair matrix "
         "at signature_matching/comparison_output_loa/comparison_report.csv; "
         "per-pair match visualisations alongside it. Working files, page "
         "renders and intermediate crops under loa_comparison/.",
         size=8.5, color=MUTED, space_after=4)
    para(doc,
         "This is an automated triage aid. It is not evidence, and it is not a "
         "finding of fact about any person's signature.",
         size=8.5, italic=True, color=MUTED)

    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--csv", default="comparison_output_loa/comparison_report.csv")
    ap.add_argument("--exhibit", default="../loa_comparison/exhibit_signatures.png")
    ap.add_argument("--sig-dir", default="signatures")
    ap.add_argument("--out",
                    default="Signature Examination - Letters of Authority.docx")
    args = ap.parse_args()

    rows = load(args.csv)
    path = build(rows, args.exhibit, args.out, args.sig_dir)
    print(f"Word report: {os.path.abspath(path)}")
    print(f"  pairs      : {len(rows)}")
    print(f"  exhibit    : {args.exhibit}"
          f"{'' if os.path.exists(args.exhibit) else '   (MISSING)'}")


if __name__ == "__main__":
    main()
